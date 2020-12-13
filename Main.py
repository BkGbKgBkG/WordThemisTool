import sys, fitz
from PyPDF2 import PdfFileReader
import numpy as np
from PIL import Image
import os, glob
import comtypes.client
import shutil
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from xlwt import Workbook, XFStyle, Borders, Pattern, Font

print("Tuỳ chọn chấm bài?")
print("1. Chấm bài theo nội dung và định dạng riêng")
print("2. Chấm bài bằng cách so sánh độ tương đồng của hai ảnh")
c = 0
while True:
    c = int(input("Nhập lựa chọn: "))
    if c == 1 or c == 2:
        break
if c == 2:
    # Xu li thu muc chua bai lam cua hoc sinh
    local = os.getcwd()
    # print(local)
    url_file = os.path.join
    # chuyen sang file pdf, can co COM- Ms Office 2010 hoac cao hon
    def chuyenSangPDF():
        wdFormatPDF = 17
        if not (os.path.exists("Temp")):os.mkdir("Temp")
        else: pass
        goal = os.path.join(os.getcwd(), "Temp")
        list_path = glob.glob("*.docx")
        for name in list_path:
            in_file = os.path.abspath(name)
            out_file = os.path.join(goal, name[:-4] + "pdf")
            if os.path.exists(out_file): continue
            else:
                word = comtypes.client.CreateObject("Word.Application")
                doc = word.Documents.Open(in_file)
                doc.SaveAs(out_file, FileFormat=wdFormatPDF)
                doc.Close()
                word.Quit()
    # Module chuyen tung trang word sang hinh anh
    def chuyenSangHinhAnh(ten_file_word):
        # chuyen file pdf sang hinh anh
        ten_file_pdf = ten_file_word[:-4] + "pdf"
        # dem so trang cua file
        pdf = PdfFileReader(open(ten_file_pdf, "rb"))
        all_pages = pdf.getNumPages()
        # Chuyen tu file pdf sang anh (png)
        doc = None
        try: doc = fitz.open(ten_file_pdf)
        except Exception as e:
            print(e)
            if doc:
                doc.close()
                exit(0)
        ten_anh = ten_file_word[:-5] + "_"
        # print (ten_anh)
        for i in range(all_pages):
            trang = doc[i]
            in4_anh = fitz.Matrix(fitz.Identity)
            in4_anh.preScale(2, 2)
            anh = trang.getPixmap(alpha=False, matrix=in4_anh)
            ten = ten_anh + str(i) + ".png"
            anh.writePNG(ten)
        return all_pages
    # Module so sanh do tuong quan giua hai buc anh chup lai trang word
    def find_corr_x_y(x, y):  # 1
        n = len(x)  # 2
        prod = []
        for xi, yi in zip(x, y):  prod.append(xi * yi)
        sum_prod_x_y = sum(prod)  # 4
        sum_x = sum(x)
        sum_y = sum(y)
        squared_sum_x = sum_x ** 2
        squared_sum_y = sum_y ** 2
        x_square = []
        for xi in x: x_square.append(xi ** 2)
        x_square_sum = sum(x_square)
        y_square = []
        for yi in y: y_square.append(yi ** 2)
        y_square_sum = sum(y_square)
        # Use formula to calculate correlation                      #5
        numerator = n * sum_prod_x_y - sum_x * sum_y
        denominator_term1 = n * x_square_sum - squared_sum_x
        denominator_term2 = n * y_square_sum - squared_sum_y
        denominator = (denominator_term1 * denominator_term2) ** 0.5
        correlation = numerator / denominator
        return correlation
    def so_sanh_anh(x, y):
        # load anh va chuyen ve kieu list
        anh_1 = Image.open(x)
        anh_2 = Image.open(y)
        data_anh_1 = np.asarray(anh_1).flatten().tolist()
        data_anh_2 = np.asarray(anh_2).flatten().tolist()
        # tinh toan do tuong dong
        corr_1_2 = find_corr_x_y(data_anh_1, data_anh_2)
        return corr_1_2
    """XU LI CAC FILE TRONG THU MUC DE DUA RA KET QUA"""
    try: chuyenSangPDF()
    except Exception:
        print("LỖI!!!!")
        os.system("pause")
        exit(0)
    # Nhap ten file mau
    cc = True
    while cc:
        file_goc = input("Nhập tên tệp đáp án: ")
        if os.path.isfile(url_file(local, file_goc)):cc = False
    # Tim kiem tat ca cac file *.docx trong thu muc
    list_file_word = glob.glob("*.docx")
    list_file_word.remove(file_goc)
    os.chdir(url_file(local, "Temp"))
    # Chuyen file mau thanh hinh anh de lam chuan so sanh
    all_pages_1 = chuyenSangHinhAnh(file_goc)
    file_goc = file_goc[:-5] + "_"
    # Nhap thang diem cho bai lam
    print("Tệp đáp án có " + str(all_pages_1) + " trang \n")
    diem_mau = [
        int(input("Nhập điểm trang thứ " + str(i + 1) + ": ")) for i in range(all_pages_1)
    ]
    print("\n")
    # Chuong trinh con tinh diem cua file bai lam
    def tinhKetQua(file_so_sanh):
        # Lay ten cua file annh (*.png) de so sanh voi anh cua file mau
        all_pages_2 = chuyenSangHinhAnh(file_so_sanh)
        file_so_sanh = file_so_sanh[:-5] + "_"
        ket_qua = 0
        for i in range(min(all_pages_1, all_pages_2)):
            ten_anh_1 = file_goc + str(i) + ".png"
            ten_anh_2 = file_so_sanh + str(i) + ".png"
            try: corr = so_sanh_anh(ten_anh_2, ten_anh_1)
            except Exception:  corr = 0 #Neu bai lam rong
            os.remove(ten_anh_2)
            # Diem cua hoc sinh duoc tinh bang do tuong dong nhan voi diem duoc quy dinh do nguoi cham nhap
            ket_qua += diem_mau[i] * corr
            ket_qua = round(ket_qua, 2)
        return ket_qua
    def in_ket_qua():
        # tao file ketqua.xls de ghi ket qua
        fnt = Font()
        fnt.name = "Times New Roman"
        fnt.height = 280
        borders = Borders()
        borders.left = Borders.THIN
        borders.right = Borders.THIN
        borders.top = Borders.THIN
        borders.bottom = Borders.THIN
        style = XFStyle()
        style.font = fnt
        style.borders = borders
        wb = Workbook()
        sheet1 = wb.add_sheet("Điểm")
        col = 0
        row = 0
        sheet1.write(row, col, "STT", style)
        sheet1.write(row, col + 1, "TÊN BÀI LÀM", style)
        sheet1.write(row, col + 2, "ĐIỂM", style)
        row += 1
        st = 1
        # Tinh ket qua cua tat ca cac file bai lam va in ket qua ra file KetQua.xls
        for name in list_file_word:
            sheet1.write(row, col, st, style)
            sheet1.write(row, col + 1, name[:-5], style)
            sheet1.write(row, col + 2, tinhKetQua(name), style)
            row += 1
            st += 1
            print(name, "___ĐÃ CHẤM XONG___")
        # Xoa bo cac file anh sau khi da xu li xong
        for name in glob.glob("*.png"):os.remove(name)
        # Luu ket qua ra file
        os.chdir(local)
        # shutil.rmtree('Temp')
        wb.save("KetQua.xls")
    in_ket_qua()
    print("KIỂM TRA FILE KetQua.xls ĐỂ XEM KẾT QUẢ")
##############

########################

#############
else:
    def lay_noi_dung(d): # lay ra chu
        chu = []
        for i in d.paragraphs: chu.append(i.text)
        return chu
    def in_dam(a): # lay chu in dam
        dam = []
        nghieng = []
        gach = []
        for i in a.paragraphs:
            for j in i.runs: 
                if j.font.bold == True: dam.append(j.text)
        return dam
    def in_nghieng(a): # lay chu in nghieng
        nghieng = []
        for i in a.paragraphs: 
            for j in i.runs: 
                if j.font.italic == True: nghieng.append(j.text)
        return nghieng    
    def gach_chan(a):# lay chu gach chan
        gach = []
        for i in a.paragraphs: 
            for j in i.runs: 
                if j.font.underline == True: gach.append(j.text)
        return gach
    def mau_sac(a):
        mau = []
        for i in a.paragraphs: 
            for j in i.runs: mau.append(j.font.color.rgb)
        return mau
    def co_chu(a):
        to = []
        for i in a.paragraphs:
            size = i.style.font.size
            to.append(size)
        return to
    def dinh_dang(a):
        sections = a.sections
        section = sections[0]
        orien = section.orientation
        left = section.left_margin
        right = section.right_margin
        top = section.top_margin
        bot = section.bottom_margin
        detail = [orien, left, right, top, bot]
        return detail
    ##### xu li file mau
    t = glob.glob("*.docx")
    while True:
        name_form = input("Nhập tên tệp đáp án: ")
        if name_form in t: break
    form = docx.Document(name_form)
    form_text = "".join(lay_noi_dung(form))
    form_bold = "".join(in_dam(form))
    form_italic = "".join(in_nghieng(form))
    form_under = "".join(gach_chan(form))
    form_color = mau_sac(form)
    form_margin = dinh_dang(form)
    form_size = co_chu(form)
    def ss_chung(a, b):# so sanh chung 
        mis = 0
        if len(a) == len(b): 
            for i in range(len(a)): 
                if a[i] != b[i]: mis = mis + 1
        else:
            mis = abs(len(a) - len(b))
            for i in range(min(len(a), len(b))): 
                if a[i] != b[i]: mis = mis + 1
        return 1 - mis / len(b)
    def so_sanh(a):
        n = 0
        # so noi dung
        noi_dung = "".join(lay_noi_dung(a))
        per_noi_dung = ss_chung(noi_dung, form_text)
        # so sanh chu in dam
        dam = "".join(in_dam(a))
        if len(dam) > 0 and len(form_bold) > 0:
            n = n + 1
            per_dam = ss_chung(dam, form_bold)
        else: per_dam = 0
        # so sanh chu in dam
        nghieng = "".join(in_nghieng(a))
        if len(nghieng) > 0 and len(form_italic) > 0:
            n = n + 1
            per_nghieng = ss_chung(nghieng, form_italic)
        else: per_nghieng = 0
        # so sanh chu in gach
        gach = "".join(gach_chan(a))
        if len(gach) > 0 and len(form_under) > 0:
            n = n + 1
            per_gach = ss_chung(gach, form_under)
        else: per_gach = 0
        # so sanh dinh dang
        per_dinh_dang = ss_chung(dinh_dang(a), form_margin)
        per_mau_sac = ss_chung(mau_sac(a), form_color)
        if n > 0:
            per_in = (per_dam + per_nghieng + per_gach) / n
            per_chung = (per_noi_dung + per_in + per_dinh_dang + per_mau_sac) / 4
        else: per_chung = (per_noi_dung + per_dinh_dang + per_mau_sac) / 3
        return per_chung
    def RUNN():
        # ghi file excel
        fnt = Font()
        fnt.name = "Times New Roman"
        fnt.height = 280
        borders = Borders()
        borders.left = Borders.THIN
        borders.right = Borders.THIN
        borders.top = Borders.THIN
        borders.bottom = Borders.THIN
        style = XFStyle()
        style.font = fnt
        style.borders = borders
        wb = Workbook()
        sheet1 = wb.add_sheet("Điểm")
        col = 0
        row = 0
        sheet1.write(row, col, "STT", style)
        sheet1.write(row, col + 1, "TÊN BÀI LÀM", style)
        sheet1.write(row, col + 2, "ĐIỂM", style)
        row += 1
        diem = []
        st = 1
        for i in t:
            if i == name_form: continue
            else:
                try:
                    t1 = docx.Document(i)
                    sheet1.write(row, col, st, style)
                    sheet1.write(row, col + 1, i[:-5], style)
                    sheet1.write(row, col + 2, round(10 * so_sanh(t1), 2), style)
                    row += 1
                    st += 1
                except Exception: pass
        wb.save("KetQua.xls")
    RUNN()
    print("ĐÃ CHẤM XONG")
    print("KIỂM TRA FILE KetQua.xls ĐỂ XEM KẾT QUẢ")
os.system("pause")
